"""
PNU-Pathfinder 해커톤 예선 개발계획서 Word 생성기
실행: python3 scripts/generate_word.py
출력: output/개발계획서_PNU-Pathfinder.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 페이지 여백 설정 ────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)

PNU_BLUE = RGBColor(0x00, 0x38, 0x76)
ACCENT   = RGBColor(0x05, 0x7A, 0x55)
BLACK    = RGBColor(0x00, 0x00, 0x00)

# ── 스타일 헬퍼 ─────────────────────────────────────────
def set_font(run, size=11, bold=False, color=BLACK, name="맑은 고딕"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)


def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=PNU_BLUE)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # 하단 테두리
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '003876')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=ACCENT)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '003876')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    for r_i, row_data in enumerate(rows):
        row = table.rows[r_i + 1]
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


# ════════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("제7회 PNU 창의융합AI해커톤")
set_font(run, size=14, color=PNU_BLUE)
p.paragraph_format.space_before = Pt(72)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("개발계획서")
set_font(run, size=28, bold=True, color=PNU_BLUE)
p.paragraph_format.space_before = Pt(24)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PNU-Pathfinder")
set_font(run, size=22, bold=True, color=ACCENT)
p.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("부산대학교 학생 맞춤형 학업·진로 자기경영 AI 플랫폼")
set_font(run, size=13, color=BLACK)
p.paragraph_format.space_before = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n트랙: 융합트랙  |  주제유형: 자유주제\n팀명: [팀명 입력]\n제출일: 2026년 5월 11일")
set_font(run, size=11, color=RGBColor(0x47, 0x55, 0x69))
p.paragraph_format.space_before = Pt(48)

doc.add_page_break()


# ════════════════════════════════════════════════════════
# 1. 프로젝트 개요
# ════════════════════════════════════════════════════════
h1("1. 프로젝트 개요")

h2("1-1. 프로젝트명")
body("PNU-Pathfinder: 부산대학교 학생 맞춤형 학업·진로 자기경영 AI 플랫폼")

h2("1-2. 개발 배경 및 필요성")
body(
    "부산대학교 학생들은 매 학기 수강신청, 재수강 여부 판단, 졸업요건 확인, 전공 이수 관리, "
    "비교과 활동, 자격증·어학 준비, 취업 준비 등을 스스로 관리해야 한다. "
    "그러나 이러한 정보는 학과 홈페이지, 학사 시스템, 비교과 시스템, 취업지원 사이트 등 "
    "여러 경로에 분산되어 있어 학생이 자신의 현황을 한눈에 파악하기 어렵다."
)
body(
    "특히 학생마다 이수 과목, 성적, 재수강 가능 여부, 남은 졸업요건, 희망 진로, 준비 수준이 "
    "모두 다르기 때문에 단순한 공지나 표준 이수체계도만으로는 개인에게 적합한 학업·진로 "
    "계획을 수립하기 어렵다."
)

h2("1-3. 프로젝트 목표")
for item in [
    "부산대학교 학생의 학업 정보와 진로 준비 정보를 하나의 서비스로 통합",
    "학생별 이수 현황과 졸업요건 분석을 통한 개인 맞춤형 수강 계획 제안",
    "재수강 필요 과목, 부족한 졸업요건, 추천 비교과 활동 AI 기반 안내",
    "진로 목표에 따른 취업 준비 로드맵 및 이력서 작성 보조 기능 제공",
    "해커톤 기간 내 실제 시연 가능한 웹 기반 MVP 구현 및 배포",
]:
    bullet(item)


# ════════════════════════════════════════════════════════
# 2. 해결하고자 하는 문제
# ════════════════════════════════════════════════════════
h1("2. 해결하고자 하는 문제")

h2("2-1. 정보 분산 문제")
body("현재 부산대학교 학생들은 아래와 같이 여러 시스템에 흩어진 정보를 각각 확인해야 한다.")
for item in [
    "학과 홈페이지: 교육과정, 이수체계도, 전공 이수 안내",
    "학사 시스템: 성적, 수강 이력, 이수 학점",
    "비교과 시스템: 비교과 프로그램, 마일리지, 활동 내역",
    "취업지원 관련 사이트: 채용 정보, 진로 프로그램",
    "외부 사이트: 자격증 일정, 어학 시험 일정, 대외활동 정보",
]:
    bullet(item)

h2("2-2. 개인 맞춤형 해석의 부재")
body("학생마다 학번·입학년도, 학과·전공 및 복수전공 여부, 이수 과목, 성적, 남은 졸업요건, 희망 진로가 모두 다르나 기존 정보 제공 방식은 개인 상황을 반영하지 못한다.")

h2("2-3. 실행 계획 수립의 어려움")
for item in [
    "다음 학기 어떤 과목을 우선적으로 들어야 하는지 추천 부족",
    "재수강 우선순위 판단 어려움",
    "진로 목표에 맞는 비교과·자격증·어학 준비 일정 설계 부족",
    "졸업 직전 이력서 및 포트폴리오 정리 부담 증가",
]:
    bullet(item)


# ════════════════════════════════════════════════════════
# 3. 핵심 기능
# ════════════════════════════════════════════════════════
h1("3. 핵심 기능")

features = [
    ("Smart Planner (지능형 수강 시뮬레이터)",
     "영역별 잔여 학점 자동 계산 및 졸업 충족률 시각화, 선수과목·졸업 필수 반영 수강 추천, "
     "재수강 우선순위 추천, 학기별 수강 시뮬레이션"),
    ("Interactive Roadmap (학과별 로드맵 시각화)",
     "정적 PDF 이수체계도를 인터랙티브 화면으로 제공. "
     "이수 과목 Checked / 수강 가능 과목 Unlock / 미이수 필수 과목 Warning 표시"),
    ("AI 챗봇 (자연어 학업 상담)",
     "학생의 이수 현황 기반 자연어 질의응답. "
     "OpenAI API 연동, 규칙 기반 추천 결과를 AI가 설명하는 구조"),
    ("What-if 시나리오 분석",
     "복수전공·전과·부전공 시 필요한 추가 이수 계획 시뮬레이션. "
     "저학년 진로·전공 의사결정 지원"),
    ("AI 이력서 자동 생성 (Career Builder)",
     "학업 이력·프로젝트·비교과 활동을 기반으로 직무 맞춤형 이력서·자기소개서 초안 생성"),
    ("졸업 자가진단 알림",
     "부족한 졸업요건, 영어 성적 기한, 비교과 포인트 충족 여부를 능동적으로 안내"),
]
for title, desc in features:
    h2(title)
    body(desc)


# ════════════════════════════════════════════════════════
# 4. AI 활용 계획
# ════════════════════════════════════════════════════════
h1("4. AI 활용 계획")

h2("4-1. 서비스 내 AI 활용")
for item in [
    "맞춤형 수강·재수강 추천 생성",
    "졸업요건·이수체계 자연어 질의응답 (OpenAI API)",
    "추천 결과를 이해하기 쉬운 자연어로 해설 (설명 가능한 AI)",
    "이력서·자기소개서 초안 자동 생성",
    "커뮤니티 팁 요약 및 핵심 인사이트 제공",
]:
    bullet(item)

h2("4-2. 개발 과정 AI 활용")
add_table(
    ["AI 도구", "활용 단계", "활용 내용"],
    [
        ["Claude (Anthropic)", "기획·설계·문서", "문제 정의, 기능 정리, 개발계획서 작성, 코드 리뷰"],
        ["ChatGPT / Gemini", "아이디어 검토", "서비스 구조 검토, 사용자 시나리오 작성"],
        ["Cursor", "개발 전 단계", "코드 생성, 디버깅, 리팩토링"],
        ["GitHub Copilot", "개발 전 단계", "함수 자동완성, 테스트 코드 생성"],
        ["Midjourney", "UI/UX 설계", "화면 시안 및 아이콘 생성"],
    ],
    col_widths=[3.5, 3.0, 8.5]
)

h2("4-3. AI 적용 방식의 특징")
body(
    "정확성이 중요한 졸업요건 및 수강 추천 영역은 규칙 기반 추천 엔진 + 생성형 AI 설명 방식으로 설계한다. "
    "실제 추천 결과는 DB와 조건 로직으로 계산하고, AI는 이를 자연어로 해설하는 역할을 맡아 "
    "신뢰성과 설명 가능성을 동시에 확보한다."
)


# ════════════════════════════════════════════════════════
# 5. AI Agent 협업 개발 방식
# ════════════════════════════════════════════════════════
h1("5. AI Agent 협업 개발 방식")

body(
    "본 프로젝트는 단순히 AI 도구를 보조적으로 활용하는 것을 넘어, "
    "Claude 멀티 에이전트 시스템과 MCP(Model Context Protocol)를 활용하여 "
    "PM·Backend·Frontend Agent가 실제로 역할을 분담하고 협업하는 방식으로 개발된다."
)

h2("5-1. 에이전트 구성")
add_table(
    ["Agent", "모델", "담당 역할", "MCP 서버"],
    [
        ["PM Agent", "Claude Sonnet 4.6",
         "기획, 문서 작성, 코드 리뷰, 진행 관리",
         "filesystem, github"],
        ["Backend Agent", "Claude Sonnet 4.6",
         "FastAPI 구현, DB 설계, 추천 로직, OpenAI 연동",
         "filesystem, postgres, github"],
        ["Frontend Agent", "Claude Sonnet 4.6",
         "React UI 구현, 컴포넌트 작성, API 연동",
         "filesystem, github, puppeteer"],
    ],
    col_widths=[3.0, 3.5, 5.5, 4.0]
)

h2("5-2. MCP 서버별 역할")
for item in [
    "filesystem — 각 Agent가 담당 폴더(backend/ 또는 frontend/)의 코드를 직접 읽고 수정",
    "postgres — Backend Agent가 PostgreSQL DB 스키마를 직접 조회하고 마이그레이션 검증",
    "github — 모든 Agent가 자신의 작업을 git commit·push하고 PR 생성",
    "puppeteer — Frontend Agent가 개발 서버에서 UI를 자동으로 스크린샷 촬영하여 렌더링 검증",
]:
    bullet(item)

h2("5-3. Agent Skills")
for item in [
    "claude-api (Backend Agent): OpenAI API 연동 패턴, 프롬프트 캐싱 최적화",
    "simplify (Backend·Frontend Agent): 작성된 코드의 중복 제거 및 품질 자동 검토",
    "security-review (PM Agent): PR 머지 전 보안 취약점 자동 스캔",
    "review (PM Agent): 코드 변경사항 전체 리뷰 및 피드백 생성",
]:
    bullet(item)

h2("5-4. 협업 흐름")
body(
    "PM Agent가 PROGRESS.md를 읽어 현재 단계를 파악한 후 태스크를 분류한다. "
    "Backend Agent와 Frontend Agent는 worktree 격리 방식으로 병렬 실행되어 충돌 없이 동시 작업한다. "
    "각 Agent의 작업은 git commit으로 자동 기록되며, PM Agent가 PR을 리뷰한 후 develop 브랜치에 병합한다."
)

h2("5-5. 발표 시 강조 포인트")
for item in [
    "AI가 코드를 '제안'하는 수준이 아니라, 실제로 코드를 작성·커밋·PR 생성까지 수행",
    "postgres MCP를 통해 Backend Agent가 실시간으로 DB 상태를 확인하며 코드 작성",
    "puppeteer MCP를 통해 Frontend Agent가 UI 렌더링을 자동 검증",
    "모든 AI 작업 이력이 GitHub 커밋 히스토리로 투명하게 추적 가능",
]:
    bullet(item)


# ════════════════════════════════════════════════════════
# 6. 기술 스택
# ════════════════════════════════════════════════════════
h1("6. 기술 스택")
add_table(
    ["구분", "기술", "역할"],
    [
        ["Frontend", "React 19 + Vite 7 + TypeScript", "UI 컴포넌트, 페이지 라우팅, API 연동"],
        ["Backend", "FastAPI (Python) + uvicorn", "REST API, 인증, 추천 로직"],
        ["Database", "PostgreSQL + SQLAlchemy 2.0", "학생 정보, 과목, 졸업요건 저장"],
        ["인증", "JWT (PyJWT) + bcrypt", "학생 로그인·회원가입"],
        ["AI", "OpenAI API (GPT-4o)", "챗봇, 이력서 생성, 자연어 설명"],
        ["배포", "Vercel (FE) + VPS/Docker (BE)", "서비스 배포 및 운영"],
        ["협업/AI도구", "GitHub + Cursor + Claude", "버전 관리, AI 코딩 보조"],
    ],
    col_widths=[3.0, 5.0, 7.0]
)


# ════════════════════════════════════════════════════════
# 7. MVP 구현 범위
# ════════════════════════════════════════════════════════
h1("7. MVP 구현 범위")

h2("7-1. 우선 구현 기능 (해커톤 기간 내)")
for item in [
    "학생 회원가입 / 로그인 (완성)",
    "학과별 졸업요건 DB 및 관리자 CRUD (완성)",
    "수강 이력 입력 및 조회",
    "졸업요건 충족률 계산 대시보드",
    "AI 챗봇 기반 학업 요약 설명 (OpenAI 연동)",
    "다음 학기 추천 과목 제안",
    "재수강 우선순위 추천",
]:
    bullet(item)

h2("7-2. 고도화 기능 (추가 구현 목표)")
for item in [
    "Interactive Roadmap (게임 스킬트리 방식)",
    "What-if 시나리오 분석",
    "AI 이력서·자기소개서 초안 생성",
    "비교과·자격증 추천 타임라인",
    "졸업 자가진단 알림",
]:
    bullet(item)


# ════════════════════════════════════════════════════════
# 7. 개발 일정
# ════════════════════════════════════════════════════════
h1("8. 개발 일정")
add_table(
    ["단계", "기간", "주요 내용"],
    [
        ["예선 준비", "2026.05.04 ~ 05.11", "개발계획서·발표자료 제출"],
        ["예선 결과", "2026.05.19", "결과 발표 및 멘토링"],
        ["1단계 개발", "2026.06.01 ~ 06.28", "수강이력 기능, AI 챗봇 연동, 기본 대시보드"],
        ["2단계 개발", "2026.07.01 ~ 07.22", "What-if, 이력서 생성, Interactive Roadmap"],
        ["중간보고회", "2026.07.23 ~ 07.24", "중간 결과 발표 및 피드백"],
        ["3단계 개발", "2026.07.25 ~ 08.26", "고도화, 테스트, 배포"],
        ["최종 발표", "2026.08.28", "농심호텔 연회장 발표·시연"],
    ],
    col_widths=[3.0, 4.5, 7.5]
)


# ════════════════════════════════════════════════════════
# 8. 팀 구성 및 역할
# ════════════════════════════════════════════════════════
h1("9. 팀 구성 및 역할 분담")
add_table(
    ["역할", "구분", "주요 담당 업무"],
    [
        ["PM / 기획", "SW전공자", "프로젝트 관리, 문서·발표자료, AI 도구 총괄, GitHub 관리"],
        ["Backend 개발", "SW전공자", "FastAPI 서버, DB 설계, 추천 로직, OpenAI 연동"],
        ["Frontend 개발", "SW전공자 또는 비전공자", "React UI 개발, 화면 설계, API 연동"],
        ["기획·데이터", "비전공자", "사용자 요구사항, 교육과정 데이터 정리, 시연 기획"],
    ],
    col_widths=[3.0, 4.0, 8.0]
)


# ════════════════════════════════════════════════════════
# 9. 기대효과 및 차별성
# ════════════════════════════════════════════════════════
h1("10. 기대효과 및 차별성")

h2("10-1. 차별성")
for item in [
    "학업·졸업·진로·이력서 준비를 하나의 플랫폼에서 통합 관리",
    "학생별 이수 현황에 따른 맞춤형 수강 로드맵 제공",
    "정형 데이터 기반 분석과 생성형 AI 설명을 결합한 신뢰성 높은 추천",
    "정적 PDF가 아닌 인터랙티브 로드맵 UI 제공",
    "복수전공·전과 시나리오 분석 등 실질적인 의사결정 지원",
]:
    bullet(item)

h2("10-2. 기대효과")
for item in [
    "학생의 수강 계획 수립 시간 단축",
    "졸업요건 누락 방지 및 학업 관리 효율 향상",
    "진로 준비의 체계화 및 비교과 참여 활성화",
    "편입생·복학생·복수전공 학생의 학업 적응 지원",
    "취업 준비 자료 정리와 이력서 작성 부담 완화",
]:
    bullet(item)


# ── 저장 ────────────────────────────────────────────────
out_dir = os.path.join(os.path.dirname(__file__), "..", "output")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "개발계획서_PNU-Pathfinder.docx")
doc.save(out_path)
print(f"✅ Word 생성 완료: {out_path}")
